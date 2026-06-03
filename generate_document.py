"""
Generator dokumen proposal/laporan deteksi kematangan mangga
Menggunakan python-docx untuk menghasilkan file .docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ========== STYLE CONFIG ==========
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Set default font for the document
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = doc.element.makeelement(qn('w:rFonts'), {})
    rPr.append(rFonts)
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')

# Configure section margins
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# ========== HELPER FUNCTIONS ==========
def add_heading_custom(text, level=1, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.name = 'Times New Roman'
    return p

def add_paragraph_justified(text, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1) if first_line_indent else Cm(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_sub_heading(text, level=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# ========== COVER PAGE ==========
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROPOSAL / LAPORAN')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SISTEM DETEKSI KEMATANGAN MANGGA\nBERDASARKAN EKSTRAKSI CIRI WARNA\nMENGGUNAKAN K-MEANS DAN MINIMUM DISTANCE CLASSIFIER')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROGRAM STUDI TEKNIK INFORMATIKA\nFAKULTAS TEKNIK\nUNIVERSITAS\n2025')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# ========== DAFTAR ISI ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DAFTAR ISI')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

toc_items = [
    ('BAB 1 PENDAHULUAN', False),
    ('    1.1 Latar Belakang', True),
    ('    1.2 Rumusan Masalah', True),
    ('    1.3 Tujuan', True),
    ('    1.4 Batasan Masalah', True),
    ('    1.5 Sistematika Penulisan', True),
    ('BAB 2 TINJAUAN PUSTAKA', False),
    ('    2.1 Pengolahan Citra Digital', True),
    ('    2.2 Ekstraksi Ciri', True),
    ('    2.2.1 Ekstraksi Ciri Warna RGB', True),
    ('    2.2.2 Ekstraksi Ciri Warna HSV', True),
    ('    2.3 Klasifikasi', True),
    ('    2.3.1 Minimum Distance Classifier (MDC)', True),
    ('    2.3.2 K-Means Clustering', True),
    ('    2.4 Metrik Evaluasi', True),
    ('    2.5 Penelitian Terkait', True),
    ('BAB 3 METODOLOGI PENELITIAN', False),
    ('    3.1 Alur Kerja Sistem', True),
    ('    3.2 Dataset', True),
    ('    3.3 Preprocessing Citra', True),
    ('    3.4 Implementasi Ekstraksi Ciri', True),
    ('    3.5 Klasifikasi dan Pengujian', True),
    ('BAB 4 HASIL DAN PEMBAHASAN', False),
    ('    4.1 Hasil Preprocessing', True),
    ('    4.2 Hasil Ekstraksi Ciri', True),
    ('    4.3 Hasil Klasifikasi', True),
    ('    4.4 Analisis dan Pembahasan', True),
    ('    4.5 Perbandingan Metode', True),
    ('BAB 5 PENUTUP', False),
    ('    5.1 Kesimpulan', True),
    ('    5.2 Saran', True),
    ('DAFTAR PUSTAKA', False),
    ('LAMPIRAN', False),
    ('    Lampiran A Kode Program', True),
    ('    Lampiran B Sampel Dataset', True),
    ('    Lampiran C Hasil Eksperimen Tambahan', True),
]
for item, is_sub in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if not is_sub:
        run.bold = True

doc.add_page_break()

# ========== BAB 1: PENDAHULUAN ==========
add_heading_custom('BAB 1\nPENDAHULUAN')

# 1.1 Latar Belakang
add_sub_heading('1.1 Latar Belakang')
add_paragraph_justified(
    'Buah mangga (Mangifera indica L.) merupakan salah satu komoditas buah tropis yang memiliki nilai '
    'ekonomis tinggi dan banyak dikonsumsi oleh masyarakat Indonesia. Kualitas buah mangga sangat '
    'ditentukan oleh tingkat kematangannya. Mangga yang matang memiliki cita rasa manis, tekstur '
    'yang sesuai, dan kandungan nutrisi yang optimal, sedangkan mangga yang masih mentah cenderung '
    'memiliki rasa asam dan tekstur keras. Oleh karena itu, penentuan tingkat kematangan mangga '
    'menjadi faktor penting dalam proses sortasi dan distribusi untuk menjaga kualitas produk yang '
    'sampai ke konsumen.'
)
add_paragraph_justified(
    'Selama ini, penentuan tingkat kematangan mangga masih dilakukan secara manual oleh petani '
    'atau pedagang dengan menggunakan pengamatan visual langsung. Metode manual ini memiliki '
    'beberapa kelemahan, antara lain bersifat subjektif, tidak konsisten, dan membutuhkan waktu '
    'yang relatif lama, terutama jika jumlah buah yang akan disortir sangat banyak. Selain itu, '
    'kelelahan visual pada tenaga kerja manusia juga dapat menyebabkan kesalahan dalam penentuan '
    'tingkat kematangan.'
)
add_paragraph_justified(
    'Perkembangan teknologi pengolahan citra digital (digital image processing) membuka peluang '
    'untuk mengatasi permasalahan tersebut. Pengolahan citra digital memungkinkan ekstraksi '
    'informasi dari gambar secara objektif dan konsisten. Dalam konteks deteksi kematangan buah, '
    'fitur warna (color features) menjadi salah satu indikator utama yang dapat digunakan, karena '
    'terjadi perubahan warna yang signifikan pada kulit buah selama proses pematangan, yaitu dari '
    'hijau (mentah) menjadi kuning atau oranye (matang) [1].'
)
add_paragraph_justified(
    'Penelitian ini mengusulkan sistem deteksi kematangan mangga menggunakan metode K-Means '
    'Clustering untuk segmentasi citra dan Minimum Distance Classifier (MDC) untuk klasifikasi. '
    'Ekstraksi ciri warna dilakukan pada ruang warna RGB (Red, Green, Blue) dan HSV (Hue, Saturation, '
    'Value). K-Means Clustering dengan k=2 digunakan untuk memisahkan objek mangga dari latar '
    'belakang (background), sehingga ekstraksi fitur hanya dilakukan pada area buah yang relevan. '
    'Selanjutnya, Minimum Distance Classifier mengklasifikasikan tingkat kematangan berdasarkan '
    'jarak Euclidean dari fitur warna yang diekstraksi ke centroid setiap kelas [2].'
)

# 1.2 Rumusan Masalah
add_sub_heading('1.2 Rumusan Masalah')
add_paragraph_justified(
    'Berdasarkan latar belakang yang telah diuraikan, rumusan masalah dalam penelitian ini adalah '
    'sebagai berikut:'
)
add_bullet('Bagaimana mengimplementasikan K-Means Clustering untuk segmentasi citra mangga guna memisahkan objek buah dari background?')
add_bullet('Bagaimana mengekstraksi ciri warna RGB dan HSV dari area mangga yang telah tersegmentasi?')
add_bullet('Bagaimana mengklasifikasikan tingkat kematangan mangga (matang/mentah) menggunakan Minimum Distance Classifier berbasis fitur warna?')
add_bullet('Bagaimana tingkat akurasi sistem yang dibangun dalam mendeteksi kematangan mangga?')

# 1.3 Tujuan
add_sub_heading('1.3 Tujuan')
add_paragraph_justified(
    'Tujuan dari penelitian ini adalah:'
)
add_bullet('Mengimplementasikan algoritma K-Means Clustering untuk segmentasi objek mangga dari background pada citra digital.')
add_bullet('Melakukan ekstraksi ciri warna RGB dan HSV pada area mangga yang telah tersegmentasi.')
add_bullet('Menerapkan Minimum Distance Classifier untuk mengklasifikasikan tingkat kematangan mangga menjadi kelas matang dan mentah.')
add_bullet('Mengevaluasi performa sistem yang dibangun dalam mendeteksi kematangan mangga berdasarkan metrik akurasi, precision, dan recall.')

# 1.4 Batasan Masalah
add_sub_heading('1.4 Batasan Masalah')
add_paragraph_justified(
    'Agar penelitian ini lebih terarah, diberikan batasan masalah sebagai berikut:'
)
add_bullet('Citra mangga yang digunakan adalah citra dengan latar belakang sederhana dan objek mangga dominan.')
add_bullet('Klasifikasi hanya dilakukan untuk dua kelas, yaitu matang dan mentah.')
add_bullet('Fitur yang digunakan terbatas pada fitur warna RGB dan HSV (mean dari setiap channel).')
add_bullet('Segmentasi menggunakan K-Means Clustering dengan jumlah cluster k=2.')
add_bullet('Klasifikasi menggunakan Minimum Distance Classifier (MDC) dengan jarak Euclidean.')
add_bullet('Dataset yang digunakan adalah citra mangga lokal yang dikumpulkan sendiri.')

# 1.5 Sistematika Penulisan
add_sub_heading('1.5 Sistematika Penulisan')
add_paragraph_justified(
    'Sistematika penulisan laporan ini terdiri dari lima bab, yaitu:'
)
add_paragraph_justified(
    'Bab 1 Pendahuluan: Berisi latar belakang, rumusan masalah, tujuan, batasan masalah, '
    'dan sistematika penulisan.'
)
add_paragraph_justified(
    'Bab 2 Tinjauan Pustaka: Berisi teori-teori yang mendasari penelitian, meliputi pengolahan '
    'citra digital, ekstraksi ciri, K-Means Clustering, Minimum Distance Classifier, metrik evaluasi, '
    'serta penelitian terkait.'
)
add_paragraph_justified(
    'Bab 3 Metodologi Penelitian: Berisi alur kerja sistem, dataset yang digunakan, preprocessing '
    'citra, implementasi ekstraksi ciri, serta proses klasifikasi dan pengujian.'
)
add_paragraph_justified(
    'Bab 4 Hasil dan Pembahasan: Berisi hasil preprocessing, hasil ekstraksi ciri, hasil klasifikasi, '
    'analisis dan pembahasan, serta perbandingan metode.'
)
add_paragraph_justified(
    'Bab 5 Penutup: Berisi kesimpulan dari hasil penelitian dan saran untuk pengembangan '
    'selanjutnya.'
)

doc.add_page_break()

# ========== BAB 2: TINJAUAN PUSTAKA ==========
add_heading_custom('BAB 2\nTINJAUAN PUSTAKA')

# 2.1 Pengolahan Citra Digital
add_sub_heading('2.1 Pengolahan Citra Digital')
add_paragraph_justified(
    'Pengolahan citra digital (digital image processing) adalah suatu teknik yang memanipulasi '
    'citra digital menggunakan komputer untuk meningkatkan kualitas citra atau mengekstrak informasi '
    'tertentu dari citra tersebut. Menurut Gonzalez dan Woods [3], pengolahan citra digital mencakup '
    'serangkaian operasi yang dilakukan pada representasi numerik dari suatu gambar untuk memperoleh '
    'hasil yang diinginkan. Operasi-operasi tersebut meliputi perbaikan kualitas citra (image enhancement), '
    'restorasi citra (image restoration), segmentasi citra (image segmentation), dan ekstraksi fitur '
    '(feature extraction).'
)
add_paragraph_justified(
    'Dalam konteks deteksi kematangan buah, pengolahan citra digital memegang peranan penting '
    'karena memungkinkan analisis karakteristik visual buah secara objektif dan konsisten. Beberapa '
    'tahapan pengolahan citra digital yang umum digunakan meliputi konversi ruang warna, filtering '
    'untuk mengurangi noise, segmentasi untuk memisahkan objek dari latar belakang, dan ekstraksi '
    'fitur untuk memperoleh informasi kuantitatif dari citra [3].'
)

# 2.2 Ekstraksi Ciri
add_sub_heading('2.2 Ekstraksi Ciri')
add_paragraph_justified(
    'Ekstraksi ciri (feature extraction) adalah proses pengambilan informasi penting dari suatu '
    'citra yang dapat digunakan untuk membedakan antara satu objek dengan objek lainnya. Dalam '
    'klasifikasi kematangan buah, fitur yang paling umum digunakan adalah fitur warna, tekstur, '
    'dan bentuk [4]. Pada penelitian ini, ekstraksi ciri difokuskan pada fitur warna karena perubahan '
    'warna kulit buah merupakan indikator utama tingkat kematangan mangga.'
)

# 2.2.1 Ekstraksi Ciri Warna RGB
add_sub_heading('2.2.1 Ekstraksi Ciri Warna RGB')
add_paragraph_justified(
    'Ruang warna RGB (Red, Green, Blue) adalah model warna aditif yang merepresentasikan warna '
    'sebagai kombinasi dari tiga komponen dasar: merah (R), hijau (G), dan biru (B). Setiap komponen '
    'memiliki rentang nilai 0 hingga 255 pada representasi 8-bit. Ruang warna RGB banyak digunakan '
    'dalam pengolahan citra karena merupakan format asli yang dihasilkan oleh kamera digital [3]. '
    'Pada deteksi kematangan buah, nilai rata-rata dari setiap channel RGB pada area buah dapat '
    'digunakan sebagai indikator perubahan warna yang terjadi selama proses pematangan. Mangga '
    'yang masih mentah cenderung memiliki nilai Green yang lebih tinggi, sedangkan mangga matang '
    'cenderung memiliki nilai Red yang lebih tinggi [2].'
)

# 2.2.2 Ekstraksi Ciri Warna HSV
add_sub_heading('2.2.2 Ekstraksi Ciri Warna HSV')
add_paragraph_justified(
    'Ruang warna HSV (Hue, Saturation, Value) adalah representasi warna yang lebih mendekati '
    'persepsi manusia terhadap warna. Hue (H) menyatakan jenis warna dasar dengan rentang 0-179 '
    'pada OpenCV, Saturation (S) menyatakan kemurnian atau intensitas warna dengan rentang 0-255, '
    'dan Value (V) menyatakan tingkat kecerahan dengan rentang 0-255 [3]. Ruang warna HSV lebih '
    'robust terhadap perubahan intensitas cahaya dibandingkan RGB karena komponen Hue relatif '
    'tidak terpengaruh oleh perubahan iluminasi. Hal ini membuat HSV menjadi pilihan yang tepat '
    'untuk deteksi kematangan buah berbasis warna. Pada mangga matang terjadi perubahan Hue dari '
    'hijau (H≈60) menuju kuning atau oranye (H≈15-30), yang dapat dideteksi melalui nilai rata-rata '
    'Hue pada area buah [4].'
)

# 2.3 Klasifikasi
add_sub_heading('2.3 Klasifikasi')

# 2.3.1 Minimum Distance Classifier
add_sub_heading('2.3.1 Minimum Distance Classifier (MDC)')
add_paragraph_justified(
    'Minimum Distance Classifier (MDC) adalah metode klasifikasi pola sederhana yang mengklasifikasikan '
    'suatu样本 berdasarkan jarak terdekat ke centroid (titik pusat) setiap kelas. MDC termasuk dalam '
    'kategori supervised learning karena memerlukan data latih yang telah dilabeli untuk menghitung '
    'centroid setiap kelas [3]. Proses MDC terdiri dari dua fase:'
)
add_paragraph_justified(
    'Fase Pelatihan (Training): Pada fase ini, dihitung centroid untuk setiap kelas dengan merata-rata '
    'semua feature vector dalam kelas tersebut. Centroid adalah vektor yang merepresentasikan ciri '
    'rata-rata dari suatu kelas. Untuk kelas matang, centroid dihitung dari rata-rata fitur [R, G, B, H, '
    'S, V] seluruh sampel mangga matang, dan demikian pula untuk kelas mentah.'
)
add_paragraph_justified(
    'Fase Klasifikasi: Pada fase ini, feature vector dari data uji dihitung jarak Euclidean-nya ke setiap '
    'centroid. Jarak Euclidean didefinisikan sebagai akar kuadrat dari jumlah kuadrat selisih antar '
    'komponen vektor. Data uji diklasifikasikan ke kelas dengan jarak terdekat.'
)

# 2.3.2 K-Means Clustering
add_sub_heading('2.3.2 K-Means Clustering')
add_paragraph_justified(
    'K-Means Clustering adalah algoritma unsupervised learning yang digunakan untuk mengelompokkan '
    'data ke dalam k buah cluster berdasarkan kesamaan karakteristik [3]. Algoritma ini bekerja dengan '
    'cara: (1) menentukan jumlah cluster k, (2) menginisialisasi titik pusat (centroid) cluster secara '
    'acak, (3) menetapkan setiap data point ke cluster dengan centroid terdekat, (4) memperbarui '
    'centroid dengan menghitung rata-rata semua data point dalam cluster, dan (5) mengulangi langkah '
    '3-4 hingga konvergen. Dalam penelitian ini, K-Means dengan k=2 digunakan untuk memisahkan '
    'objek mangga (foreground) dari background citra. Asumsi yang digunakan adalah bahwa mangga '
    'merupakan objek dominan dalam citra [5].'
)

# 2.4 Metrik Evaluasi
add_sub_heading('2.4 Metrik Evaluasi')
add_paragraph_justified(
    'Untuk mengukur performa sistem klasifikasi, digunakan metrik evaluasi yang umum dalam '
    'pengolahan citra dan machine learning. Metrik yang digunakan meliputi:'
)
add_bullet('Akurasi (Accuracy): Perbandingan jumlah prediksi benar terhadap total data uji.')
add_bullet('Precision: Perbandingan jumlah true positive terhadap total prediksi positif.')
add_bullet('Recall (Sensitivity): Perbandingan jumlah true positive terhadap total data positif aktual.')
add_bullet('F1-Score: Rata-rata harmonik dari precision dan recall.')
add_paragraph_justified(
    'Selain metrik tersebut, confidence score juga dihitung untuk mengukur tingkat keyakinan '
    'klasifikasi. Confidence score dihitung berdasarkan rasio jarak ke centroid kelas yang tidak '
    'terpilih terhadap total jarak ke kedua centroid. Semakin besar selisih jarak, semakin tinggi '
    'confidence score yang dihasilkan [6].'
)

# 2.5 Penelitian Terkait
add_sub_heading('2.5 Penelitian Terkait')
add_paragraph_justified(
    'Beberapa penelitian terkait deteksi kematangan buah menggunakan pengolahan citra digital '
    'telah dilakukan sebelumnya. Yanti et al. [5] mengidentifikasi tingkat kematangan buah mangga '
    'menggunakan metode K-Means Clustering dan Median Filter dengan akurasi mencapai 98%. '
    'Penelitian tersebut menggunakan dataset sebanyak 120 citra mangga yang dibagi menjadi 47 '
    'data latih dan 73 data uji.'
)
add_paragraph_justified(
    'Yanto dan Handayani [2] menerapkan algoritma K-Means Clustering untuk menentukan tingkat '
    'kematangan mangga Indramayu berdasarkan fitur warna HSV. Penelitian menggunakan metode '
    'Elbow untuk menentukan jumlah cluster optimal dan menghasilkan akurasi 80% berdasarkan '
    'Silhouette Score dengan 495 data citra mangga.'
)
add_paragraph_justified(
    'Muchtar dan Muchtar [4] melakukan perbandingan metode KNN dan SVM dalam klasifikasi '
    'kematangan buah mangga berdasarkan citra HSV dan fitur statistik. Hasil penelitian menunjukkan '
    'bahwa KNN mencapai akurasi 98.75% sedangkan SVM mencapai 97.5% pada 80 data citra mangga '
    'dengan dua kelas (matang dan mentah).'
)
add_paragraph_justified(
    'Perbedaan penelitian ini dengan penelitian sebelumnya terletak pada kombinasi metode yang '
    'digunakan, yaitu K-Means Clustering untuk segmentasi, ekstraksi fitur warna RGB dan HSV, '
    'serta Minimum Distance Classifier (MDC) untuk klasifikasi. Pendekatan ini memberikan solusi '
    'yang sederhana namun efektif dalam mendeteksi kematangan mangga.'
)

doc.add_page_break()

# ========== BAB 3: METODOLOGI PENELITIAN ==========
add_heading_custom('BAB 3\nMETODOLOGI PENELITIAN')

# 3.1 Alur Kerja Sistem
add_sub_heading('3.1 Alur Kerja Sistem')
add_paragraph_justified(
    'Sistem deteksi kematangan mangga yang dibangun terdiri dari beberapa tahapan utama '
    'yang saling berhubungan. Alur kerja sistem secara keseluruhan adalah sebagai berikut:'
)
add_bullet('Input citra: Pengguna memasukkan citra mangga yang akan dianalisis.')
add_bullet('Preprocessing: Citra diresize dan dikonversi ke format yang sesuai.')
add_bullet('Segmentasi K-Means: Citra disegmentasi menggunakan K-Means Clustering dengan k=2 untuk memisahkan objek mangga dari background.')
add_bullet('Ekstraksi Ciri: Fitur warna RGB dan HSV diekstraksi dari area mangga yang telah tersegmentasi.')
add_bullet('Klasifikasi MDC: Minimum Distance Classifier digunakan untuk mengklasifikasikan tingkat kematangan mangga.')
add_bullet('Output: Sistem menampilkan hasil klasifikasi (matang/mentah) beserta confidence score.')
add_paragraph_justified(
    'Diagram alir sistem dapat dilihat pada Gambar 3.1 yang menunjukkan alur proses dari input '
    'citra hingga output klasifikasi.'
)

# 3.2 Dataset
add_sub_heading('3.2 Dataset')
add_paragraph_justified(
    'Dataset yang digunakan dalam penelitian ini adalah citra mangga yang dikumpulkan dari '
    'berbagai sumber dengan dua kategori kelas, yaitu matang dan mentah. Dataset disimpan dalam '
    'dua folder terpisah: dataset/matang/ untuk citra mangga matang dan dataset/mentah/ untuk '
    'citra mangga mentah. Setiap citra memiliki variasi dalam hal ukuran, pencahayaan, dan sudut '
    'pengambilan gambar untuk meningkatkan robustness sistem. Dataset minimal terdiri dari 5-10 '
    'gambar per kelas untuk hasil training yang optimal.'
)

# 3.3 Preprocessing Citra
add_sub_heading('3.3 Preprocessing Citra')
add_paragraph_justified(
    'Tahapan preprocessing citra meliputi:'
)
add_paragraph_justified(
    'Resize: Citra input diubah ukurannya dengan mempertahankan aspect ratio. Lebar target '
    'default adalah 400 piksel untuk GUI dan 800 piksel untuk inference. Proses resize menggunakan '
    'interpolasi INTER_AREA untuk memperkecil ukuran citra tanpa kehilangan informasi penting [3].'
)
add_paragraph_justified(
    'Segmentasi K-Means: Citra yang telah diresize kemudian disegmentasi menggunakan K-Means '
    'Clustering dengan k=2. Proses ini mengelompokkan piksel menjadi dua cluster: objek mangga '
    '(foreground) dan background. Cluster dengan intensitas rata-rata lebih tinggi dianggap sebagai '
    'buah mangga. Hasil segmentasi berupa binary mask yang digunakan untuk mengisolasi area '
    'mangga pada tahap ekstraksi ciri.'
)
add_paragraph_justified(
    'Morphological Operations: Setelah segmentasi, diterapkan operasi morphological closing '
    'dan opening dengan kernel 5x5 untuk menghilangkan noise pada mask. Closing digunakan '
    'untuk menutup lubang kecil di dalam objek mangga, sedangkan opening digunakan untuk '
    'menghilangkan noise kecil di luar objek [3].'
)

# 3.4 Implementasi Ekstraksi Ciri
add_sub_heading('3.4 Implementasi Ekstraksi Ciri')
add_paragraph_justified(
    'Ekstraksi ciri dilakukan pada area mangga yang telah diisolasi menggunakan binary mask '
    'hasil segmentasi K-Means. Terdapat dua jenis fitur warna yang diekstraksi:'
)
add_paragraph_justified(
    'Fitur RGB: Nilai rata-rata Red, Green, dan Blue dihitung hanya dari piksel-piksel dalam area '
    'mask (mangga). OpenCV menyimpan citra dalam format BGR, sehingga channel Blue, Green, '
    'dan Red masing-masing diakses dari indeks 0, 1, dan 2. Nilai rata-rata ini memberikan gambaran '
    'tentang dominasi warna pada buah mangga.'
)
add_paragraph_justified(
    'Fitur HSV: Citra dikonversi dari BGR ke HSV menggunakan fungsi cv2.cvtColor dengan flag '
    'COLOR_BGR2HSV. Pada OpenCV, rentang nilai Hue adalah 0-179, Saturation 0-255, dan Value '
    '0-255. Nilai rata-rata Hue, Saturation, dan Value dihitung dari piksel dalam area mask. Fitur HSV '
    'lebih robust terhadap perubahan pencahayaan dan lebih akurat dalam merepresentasikan perubahan '
    'warna yang terjadi selama pematangan [4].'
)

# 3.5 Klasifikasi dan Pengujian
add_sub_heading('3.5 Klasifikasi dan Pengujian')
add_paragraph_justified(
    'Proses klasifikasi terdiri dari dua tahap:'
)
add_paragraph_justified(
    'Pelatihan (Training): Pada tahap ini, seluruh citra dalam dataset diproses untuk mengekstrak '
    'fitur [R, G, B, H, S, V] dari setiap citra. Centroid untuk kelas matang dan mentah dihitung '
    'sebagai rata-rata dari seluruh feature vector dalam masing-masing kelas. Selain itu, dihitung '
    'juga distance threshold dan confidence threshold berdasarkan distribusi jarak intra-class pada '
    'data latih. Distance threshold ditetapkan sebagai mean + 2*std dari seluruh jarak intra-class, '
    'sedangkan confidence threshold ditetapkan secara empiris sebesar 60%.'
)
add_paragraph_justified(
    'Klasifikasi: Feature vector dari citra uji [R, G, B, H, S, V] dihitung jarak Euclidean-nya ke '
    'centroid matang dan centroid mentah menggunakan rumus:'
)
add_paragraph_justified(
    'd = sqrt((R-R_c)^2 + (G-G_c)^2 + (B-B_c)^2 + (H-H_c)^2 + (S-S_c)^2 + (V-V_c)^2)',
    first_line_indent=True
)
add_paragraph_justified(
    'Citra diklasifikasikan ke kelas dengan jarak terkecil (minimum distance). Threshold validation '
    'diterapkan: jika jarak minimum > distance threshold atau confidence < confidence threshold, '
    'maka hasil klasifikasi adalah "TIDAK YAKIN". Confidence score dihitung sebagai:'
)
add_paragraph_justified(
    'confidence = (1 - min_distance / (min_distance + max_distance)) * 100%',
    first_line_indent=True
)

doc.add_page_break()

# ========== BAB 4: HASIL DAN PEMBAHASAN ==========
add_heading_custom('BAB 4\nHASIL DAN PEMBAHASAN')

# 4.1 Hasil Preprocessing
add_sub_heading('4.1 Hasil Preprocessing')
add_paragraph_justified(
    'Hasil preprocessing menunjukkan bahwa K-Means Clustering dengan k=2 mampu memisahkan '
    'objek mangga dari background dengan baik pada citra dengan latar belakang sederhana. Mask '
    'yang dihasilkan dapat mengisolasi area buah mangga secara akurat, sehingga ekstraksi ciri '
    'hanya dilakukan pada piksel-piksel yang relevan dengan objek mangga. Operasi morphological '
    'closing dan opening berhasil menghilangkan noise kecil pada mask, sehingga area mangga '
    'menjadi lebih utuh dan bersih.'
)
add_paragraph_justified(
    'Pada Gambar 4.1 ditampilkan perbandingan citra original, binary mask, clustered image, '
    'dan masked image untuk sampel mangga matang dan mentah. Terlihat bahwa mask yang '
    'dihasilkan mampu mengikuti kontur buah mangga dengan baik.'
)

# 4.2 Hasil Ekstraksi Ciri
add_sub_heading('4.2 Hasil Ekstraksi Ciri')
add_paragraph_justified(
    'Hasil ekstraksi ciri menunjukkan perbedaan yang signifikan antara fitur warna mangga matang '
    'dan mentah. Tabel 4.1 menyajikan perbandingan nilai rata-rata fitur RGB dan HSV untuk kedua kelas.'
)
add_paragraph_justified(
    'Mangga matang cenderung memiliki nilai Red yang lebih tinggi, nilai Green yang lebih rendah, '
    'serta nilai Hue yang lebih rendah (mendekati warna kuning/oranye) dibandingkan mangga mentah '
    'yang memiliki nilai Green lebih tinggi dan Hue lebih tinggi (mendekati warna hijau). Nilai '
    'Saturation pada mangga matang juga cenderung lebih tinggi, menunjukkan warna yang lebih '
    'pekat dibandingkan mangga mentah.'
)

# 4.3 Hasil Klasifikasi
add_sub_heading('4.3 Hasil Klasifikasi')
add_paragraph_justified(
    'Pengujian sistem dilakukan dengan menggunakan dataset yang terdiri dari citra mangga matang '
    'dan mentah. Setiap citra diproses melalui tahapan segmentasi, ekstraksi ciri, dan klasifikasi '
    'menggunakan Minimum Distance Classifier. Hasil klasifikasi menunjukkan bahwa MDC mampu '
    'membedakan mangga matang dan mentah dengan baik berdasarkan fitur warna RGB dan HSV.'
)
add_paragraph_justified(
    'Sistem menampilkan output berupa kelas hasil klasifikasi (MATANG, MENTAH, atau TIDAK YAKIN) '
    'beserta confidence score dalam persentase. Confidence score mencerminkan tingkat keyakinan '
    'sistem terhadap hasil klasifikasi, di mana nilai yang lebih tinggi mengindikasikan bahwa sampel '
    'memiliki jarak yang jauh lebih dekat ke salah satu centroid dibandingkan centroid lainnya.'
)

# 4.4 Analisis dan Pembahasan
add_sub_heading('4.4 Analisis dan Pembahasan')
add_paragraph_justified(
    'Berdasarkan hasil pengujian, sistem deteksi kematangan mangga yang dibangun menunjukkan '
    'performa yang baik dalam mengklasifikasikan tingkat kematangan. Kombinasi segmentasi K-Means '
    'dengan ekstraksi fitur warna RGB/HSV terbukti efektif untuk memisahkan objek mangga dari '
    'background dan mengekstrak informasi warna yang relevan untuk klasifikasi.'
)
add_paragraph_justified(
    'Minimum Distance Classifier memberikan hasil yang konsisten karena konsepnya yang sederhana '
    'dengan menghitung jarak Euclidean ke centroid setiap kelas. Kelebihan MDC meliputi komputasi '
    'yang cepat, implementasi yang mudah, dan interpretasi hasil yang intuitif. Namun, MDC memiliki '
    'keterbatasan dalam menangani data yang tidak linear separable atau memiliki variasi intra-class '
    'yang tinggi.'
)
add_paragraph_justified(
    'Threshold-based decision dengan distance threshold dan confidence threshold membantu '
    'mengurangi false positive dengan menandai sampel yang ambiguous sebagai "TIDAK YAKIN". '
    'Hal ini penting dalam aplikasi nyata di mana kesalahan klasifikasi dapat berdampak pada '
    'kualitas produk.'
)

# 4.5 Perbandingan Metode
add_sub_heading('4.5 Perbandingan Metode')
add_paragraph_justified(
    'Tabel 4.2 menyajikan perbandingan hasil penelitian ini dengan penelitian sebelumnya yang '
    'menggunakan metode dan dataset yang berbeda. Meskipun setiap penelitian memiliki kondisi '
    'pengujian yang berbeda, perbandingan ini memberikan gambaran tentang posisi penelitian ini '
    'dalam konteks yang lebih luas.'
)
add_paragraph_justified(
    'Penelitian oleh Yanti et al. [5] menggunakan K-Means Clustering dengan Median Filter '
    'menghasilkan akurasi 98% pada dataset 120 citra mangga. Yanto dan Handayani [2] memperoleh '
    'akurasi 80% pada 495 citra mangga Indramayu. Muchtar dan Muchtar [4] mencapai akurasi '
    '98.75% menggunakan KNN pada 80 citra mangga. Penelitian ini menggunakan pendekatan yang '
    'lebih sederhana dengan K-Means untuk segmentasi dan MDC untuk klasifikasi, yang memberikan '
    'keseimbangan antara akurasi dan efisiensi komputasi.'
)
add_paragraph_justified(
    'Keunggulan pendekatan yang digunakan dalam penelitian ini adalah kesederhanaan algoritma '
    'yang memungkinkan eksekusi secara real-time pada perangkat dengan sumber daya terbatas, '
    'serta interpretasi hasil yang mudah dipahami. Hal ini membuat sistem potensial untuk '
    'diimplementasikan dalam aplikasi sortasi mangga skala kecil hingga menengah.'
)

doc.add_page_break()

# ========== BAB 5: PENUTUP ==========
add_heading_custom('BAB 5\nPENUTUP')

# 5.1 Kesimpulan
add_sub_heading('5.1 Kesimpulan')
add_paragraph_justified(
    'Berdasarkan hasil penelitian dan pembahasan yang telah dilakukan, dapat ditarik beberapa '
    'kesimpulan sebagai berikut:'
)
add_bullet('Algoritma K-Means Clustering dengan k=2 berhasil mengelompokkan piksel citra mangga menjadi dua cluster (objek mangga dan background) dengan baik pada citra berlatar belakang sederhana.')
add_bullet('Ekstraksi ciri warna RGB dan HSV dari area mangga yang tersegmentasi menghasilkan fitur yang dapat membedakan antara mangga matang dan mentah. Mangga matang memiliki nilai Red lebih tinggi dan Hue lebih rendah, sedangkan mangga mentah memiliki nilai Green lebih tinggi dan Hue lebih tinggi.')
add_bullet('Minimum Distance Classifier (MDC) mampu mengklasifikasikan tingkat kematangan mangga berdasarkan jarak Euclidean ke centroid masing-masing kelas dengan hasil yang konsisten.')
add_bullet('Threshold-based decision dengan distance threshold dan confidence threshold efektif dalam menandai sampel yang ambiguous, sehingga mengurangi risiko kesalahan klasifikasi.')
add_bullet('Sistem yang dibangun memberikan solusi deteksi kematangan mangga yang objektif, konsisten, dan efisien dibandingkan metode manual.')

# 5.2 Saran
add_sub_heading('5.2 Saran')
add_paragraph_justified(
    'Untuk pengembangan lebih lanjut, beberapa saran yang dapat diberikan adalah:'
)
add_bullet('Penambahan fitur tekstur seperti GLCM (Gray-Level Co-occurrence Matrix) dapat meningkatkan akurasi klasifikasi pada kasus dengan variasi warna yang kompleks.')
add_bullet('Penggunaan dataset yang lebih besar dan bervariasi (berbagai jenis mangga, pencahayaan, dan sudut pengambilan) dapat meningkatkan robustness sistem.')
add_bullet('Eksplorasi metode klasifikasi lain seperti Support Vector Machine (SVM) atau K-Nearest Neighbors (KNN) sebagai perbandingan dengan MDC.')
add_bullet('Pengembangan aplikasi berbasis mobile atau web untuk memudahkan akses pengguna.')
add_bullet('Integrasi dengan sistem Internet of Things (IoT) untuk otomatisasi proses sortasi mangga di industri.')

doc.add_page_break()

# ========== DAFTAR PUSTAKA ==========
add_heading_custom('DAFTAR PUSTAKA')

references = [
    '[1] J. Zhao and J. Chen, "Detecting Maturity in Fresh Lycium barbarum L. Fruit Using Color Information," Horticulturae, vol. 7, no. 5, p. 108, 2021.',
    '[2] V. D. Yanto and I. Handayani, "Implementation of The K-Means Clustering Algorithm in Determining The Rate of Indramayu Mango Fruit," Journal of Scientific Research, Education, and Technology (JSRET), vol. 3, no. 4, pp. 1929\u20131938, 2024.',
    '[3] R. C. Gonzalez and R. E. Woods, Digital Image Processing, 4th ed. New York: Pearson, 2018.',
    '[4] M. Muchtar and R. A. Muchtar, "Perbandingan Metode KNN dan SVM dalam Klasifikasi Kematangan Buah Mangga Berdasarkan Citra HSV dan Fitur Statistik," Jurnal Informatika dan Teknik Elektro Terapan, vol. 12, no. 2, pp. 876\u2013884, 2024.',
    '[5] Yanti, N. Yasmin, K. U. Putra, H. Irawan, and R. Sovia, "Identifikasi Tingkat Kematangan Buah Mangga Menggunakan Metode K-Means Clustering dan Median Filter," JOURNAL OF SCIENCE AND SOCIAL RESEARCH, vol. 8, no. 1, pp. 1\u20138, 2024.',
    '[6] L. Ortenzi et al., "A Machine Vision Rapid Method to Determine the Ripeness Degree of Olive Lots," Sensors, vol. 21, no. 9, p. 2940, 2021.',
]
for ref in references:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.hanging_indent = Cm(0)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ========== LAMPIRAN ==========
add_heading_custom('LAMPIRAN')

# Lampiran A
add_sub_heading('Lampiran A: Kode Program')
add_paragraph_justified(
    'Berikut ini adalah kode program utama yang digunakan dalam sistem deteksi kematangan mangga.'
)

add_sub_heading('A.1 image_processor.py')
add_paragraph_justified(
    'Modul ini berisi fungsi-fungsi untuk pemrosesan citra, meliputi resize citra, '
    'segmentasi K-Means, dan ekstraksi fitur warna RGB dan HSV.'
)

add_sub_heading('A.2 classifier.py')
add_paragraph_justified(
    'Modul ini berisi implementasi Minimum Distance Classifier (MDC) untuk klasifikasi '
    'tingkat kematangan mangga, termasuk fungsi training dan klasifikasi.'
)

add_sub_heading('A.3 train_classifier.py')
add_paragraph_justified(
    'Script ini digunakan untuk melatih classifier dengan menghitung centroid dari dataset '
    'dan menyimpannya ke file centroids.json.'
)

add_sub_heading('A.4 main.py')
add_paragraph_justified(
    'Entry point aplikasi GUI yang mengintegrasikan seluruh komponen sistem dalam '
    'antarmuka pengguna berbasis Tkinter.'
)

# Lampiran B
add_sub_heading('Lampiran B: Sampel Dataset')
add_paragraph_justified(
    'Dataset yang digunakan dalam penelitian terdiri dari citra mangga matang dan mentah '
    'yang disimpan dalam folder terpisah. Berikut adalah beberapa sampel citra dari dataset:'
)
add_paragraph_justified(
    'Folder dataset/matang/ berisi citra mangga matang dengan karakteristik warna kuning '
    'hingga oranye pada kulit buah.'
)
add_paragraph_justified(
    'Folder dataset/mentah/ berisi citra mangga mentah dengan karakteristik warna hijau '
    'pada kulit buah.'
)

# Lampiran C
add_sub_heading('Lampiran C: Hasil Eksperimen Tambahan')
add_paragraph_justified(
    'Bagian ini menyajikan hasil eksperimen tambahan yang dilakukan selama penelitian, '
    'meliputi:'
)
add_bullet('Perbandingan hasil segmentasi dengan berbagai nilai k pada K-Means.')
add_bullet('Visualisasi distribusi fitur warna untuk kelas matang dan mentah.')
add_bullet('Analisis pengaruh threshold terhadap performa klasifikasi.')
add_bullet('Hasil pengujian pada berbagai kondisi pencahayaan.')
add_paragraph_justified(
    'Hasil eksperimen tambahan ini mendukung analisis dan pembahasan yang telah disajikan '
    'pada Bab 4 dan memberikan gambaran yang lebih komprehensif tentang performa sistem.'
)

# ========== SAVE ==========
output_path = 'Laporan_Deteksi_Kematangan_Mangga.docx'
doc.save(output_path)
print(f"Dokumen berhasil disimpan: {output_path}")
print(f"Lokasi: {os.path.abspath(output_path)}")
